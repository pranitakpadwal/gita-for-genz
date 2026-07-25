#!/usr/bin/env python3
"""Post-processing step: inserts a small eyebrow tagline line above the
title on the auto-generated title page, for both the epub and docx
outputs. Pandoc's built-in title-block only supports title/subtitle/
author, so this patches the already-built files rather than fighting
pandoc's template.
"""
import sys
import zipfile
import shutil

TAGLINE = "A Gita for Gen Z"
POWER_LINE = "Timeless wisdom. Modern battles. A conversation that changes you."


def patch_epub(path):
    tmp_path = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        title_page_name = next(n for n in names if n.endswith("title_page.xhtml"))
        nav_name = next(n for n in names if n.endswith("nav.xhtml"))
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == title_page_name:
                    text = data.decode("utf-8")
                    text = text.replace(
                        '<h1 class="title">',
                        f'<p class="tagline">{TAGLINE}</p>\n  <h1 class="title">',
                    )
                    text = text.replace(
                        '<p class="author">',
                        f'<p class="power-line">{POWER_LINE}</p>\n  <p class="author">',
                    )
                    data = text.encode("utf-8")
                if item.filename == nav_name:
                    # Without --toc, pandoc places nav.xhtml at the EPUB
                    # root but still emits "../media/..." paths meant for
                    # files one level down (text/*.xhtml), which resolve
                    # outside the EPUB entirely. nav.xhtml is at the same
                    # level as media/, so the "../" needs to go.
                    text = data.decode("utf-8").replace("../media/", "media/")
                    data = text.encode("utf-8")
                # mimetype must stay stored, uncompressed, first entry
                compress_type = zipfile.ZIP_STORED if item.filename == "mimetype" else zipfile.ZIP_DEFLATED
                zout.writestr(item, data, compress_type=compress_type)
    shutil.move(tmp_path, path)
    print(f"Tagline added + nav.xhtml media paths fixed -> {path}")


def _insert_styled_paragraph(doc, anchor_p, text, before=True, italic=True, size_factor=0.7):
    import copy
    from docx.text.paragraph import Paragraph

    new_p_elem = copy.deepcopy(anchor_p._p)
    if before:
        anchor_p._p.addprevious(new_p_elem)
    else:
        anchor_p._p.addnext(new_p_elem)

    new_p = Paragraph(new_p_elem, anchor_p._parent)
    new_p.style = doc.styles["Subtitle"]
    for run in list(new_p.runs):
        run.text = ""
    if new_p.runs:
        new_p.runs[0].text = text
    else:
        new_p.add_run(text)
    for run in new_p.runs:
        run.italic = italic
        if run.font.size:
            run.font.size = int(run.font.size * size_factor)
    return new_p


def patch_docx(path):
    from docx import Document

    doc = Document(path)
    title_idx = next(
        i for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == "Title"
    )
    _insert_styled_paragraph(doc, doc.paragraphs[title_idx], TAGLINE, before=True)

    author_idx = next(
        i for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == "Author"
    )
    _insert_styled_paragraph(doc, doc.paragraphs[author_idx], POWER_LINE, before=True, size_factor=0.75)

    doc.save(path)
    print(f"Tagline + power line added to docx title page -> {path}")


if __name__ == "__main__":
    kind, path = sys.argv[1], sys.argv[2]
    if kind == "epub":
        patch_epub(path)
    elif kind == "docx":
        patch_docx(path)
    else:
        sys.exit(f"unknown kind {kind}")
